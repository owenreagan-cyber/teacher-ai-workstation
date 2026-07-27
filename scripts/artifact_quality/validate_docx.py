from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Inches

from .models import CheckStatus, PLACEHOLDER_PATTERNS, PreflightReport
from .profiles import ArtifactProfile


def _section_size_label(section) -> str:
    width = section.page_width.inches if section.page_width is not None else 0
    height = section.page_height.inches if section.page_height is not None else 0
    return f"{width:.2f}x{height:.2f}in"


def validate_docx(path: Path, profile: ArtifactProfile, report: PreflightReport) -> None:
    if not path.is_file():
        report.add(CheckStatus.FAIL, "DOCX input file does not exist")
        return

    try:
        document = Document(str(path))
    except Exception as exc:  # noqa: BLE001
        report.add(CheckStatus.FAIL, "DOCX file is not readable", details=str(exc))
        return

    report.add(CheckStatus.PASS, "DOCX opens successfully")
    sections = document.sections
    if not sections:
        report.add(CheckStatus.FAIL, "DOCX has no sections")
        return

    report.add(CheckStatus.PASS, f"{len(sections)} section{'s' if len(sections) != 1 else ''} detected")
    if len(sections) > 1:
        report.add(CheckStatus.WARN, "Multiple DOCX sections detected", details="Verify page layout consistency before export.")

    expected_w = profile.paper.width_points / 72.0
    expected_h = profile.paper.height_points / 72.0
    a4_detected = False
    letter_ok = True
    for index, section in enumerate(sections, start=1):
        width = section.page_width.inches if section.page_width is not None else 0
        height = section.page_height.inches if section.page_height is not None else 0
        if abs(width - 8.27) < 0.15 and abs(height - 11.69) < 0.15:
            a4_detected = True
        if not (abs(width - expected_w) < 0.15 and abs(height - expected_h) < 0.15):
            letter_ok = False
        if section.orientation != WD_ORIENT.PORTRAIT and profile.paper.orientation == "portrait":
            report.add(CheckStatus.WARN, f"Section {index} is not portrait", details=_section_size_label(section))

        top = section.top_margin.inches if section.top_margin is not None else None
        bottom = section.bottom_margin.inches if section.bottom_margin is not None else None
        left = section.left_margin.inches if section.left_margin is not None else None
        right = section.right_margin.inches if section.right_margin is not None else None
        if top is not None and abs(top - profile.margins.top_inches) > 0.20:
            report.add(CheckStatus.WARN, f"Section {index} top margin differs from profile", details=f"{top:.2f}\" vs {profile.margins.top_inches:.2f}\"")
        if bottom is not None and abs(bottom - profile.margins.bottom_inches) > 0.20:
            report.add(CheckStatus.WARN, f"Section {index} bottom margin differs from profile", details=f"{bottom:.2f}\" vs {profile.margins.bottom_inches:.2f}\"")
        if left is not None and abs(left - profile.margins.left_inches) > 0.20:
            report.add(CheckStatus.WARN, f"Section {index} left margin differs from profile", details=f"{left:.2f}\" vs {profile.margins.left_inches:.2f}\"")
        if right is not None and abs(right - profile.margins.right_inches) > 0.20:
            report.add(CheckStatus.WARN, f"Section {index} right margin differs from profile", details=f"{right:.2f}\" vs {profile.margins.right_inches:.2f}\"")

    if a4_detected:
        report.add(CheckStatus.FAIL, "DOCX section configured as A4 instead of Letter")
    elif letter_ok:
        report.add(CheckStatus.PASS, "DOCX section page size is US Letter")

    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    placeholders = [pattern for pattern in PLACEHOLDER_PATTERNS if pattern.upper() in text.upper()]
    if placeholders:
        report.add(CheckStatus.FAIL, "Unresolved placeholders detected in DOCX", details=", ".join(placeholders))
    else:
        report.add(CheckStatus.PASS, "No unresolved placeholders detected in DOCX")

    blank_run = 0
    for paragraph in document.paragraphs:
        if not paragraph.text.strip():
            blank_run += 1
            if blank_run >= 3:
                report.add(CheckStatus.WARN, "Consecutive blank paragraphs used as spacing", details="Prefer explicit spacing styles.")
                break
        else:
            blank_run = 0

    fonts = sorted({run.font.name for paragraph in document.paragraphs for run in paragraph.runs if run.font.name})
    if fonts:
        report.add(CheckStatus.PASS, f"DOCX font inventory captured ({len(fonts)} fonts)")
    else:
        report.add(CheckStatus.WARN, "DOCX font inventory unavailable")

    report.add(
        CheckStatus.WARN,
        "Final pagination not validated from DOCX alone",
        details="Export to PDF and run PDF preflight for authoritative print proof.",
    )
